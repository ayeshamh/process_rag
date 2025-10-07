from typing import Iterator
from graphrag.document import Document

class DOCXLoader:
    """
    Load DOCX (Microsoft Word) files
    """

    def __init__(self, path: str) -> None:
        """
        Initialize loader

        Args:
            path (str): path to DOCX file.
        """
        try:
            import docx
        except ImportError:
            raise ImportError(
                "python-docx package not found, please install it with "
                "`pip install python-docx`"
            )

        self.path = path

    def load(self) -> Iterator[Document]:
        """
        Load DOCX and extract text

        Returns:
            Iterator[Document]: document iterator
        """
        from docx import Document as DocxDocument  # pylint: disable=import-outside-toplevel

        doc = DocxDocument(self.path)
        
        # Extract text from paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Extract text from tables
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text.strip()]
                if row_text:
                    table_text.append(" | ".join(row_text))
        
        # Combine paragraph and table content
        content = "\n".join(paragraphs + table_text)
        
        yield Document(content, self.path) 